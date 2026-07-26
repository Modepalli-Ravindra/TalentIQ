import io
import re
import logging
from typing import Optional

logger = logging.getLogger("talentiq.resume_parser")


def parse_pdf(file_content: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            text_parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            text = "\n".join(text_parts)
            if not text.strip():
                logger.warning("PDF extraction returned empty text")
                return ""
            return _clean_text(text)
    except ImportError:
        logger.error("pdfplumber not installed")
        raise RuntimeError("PDF parsing requires pdfplumber. Install with: pip install pdfplumber")
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
        if not text.strip():
            logger.warning("DOCX extraction returned empty text")
            return ""
        return _clean_text(text)
    except ImportError:
        logger.error("python-docx not installed")
        raise RuntimeError("DOCX parsing requires python-docx. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_txt(file_content: bytes) -> str:
    try:
        for encoding in ["utf-8", "latin-1", "cp1252", "ascii"]:
            try:
                text = file_content.decode(encoding)
                return _clean_text(text)
            except (UnicodeDecodeError, LookupError):
                continue
        raise ValueError("Unable to decode text file with any supported encoding")
    except Exception as e:
        logger.error(f"TXT parsing failed: {e}")
        raise ValueError(f"Failed to parse text file: {str(e)}")


def parse_resume_file(file_content: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return parse_pdf(file_content)
    elif ext in ("docx", "doc"):
        return parse_docx(file_content)
    elif ext == "txt":
        return parse_txt(file_content)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT")


def _clean_text(text: str) -> str:
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"javascript:", "", text, flags=re.IGNORECASE)
    text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"on\w+\s*=\s*['\"][^'\"]*['\"]", "", text, flags=re.IGNORECASE)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    return text


def extract_sections(text: str) -> dict:
    section_patterns = {
        "contact": r"(?:contact|info|details)\s*\n(.*?)(?=\n\s*\n|\n(?:experience|education|skills|projects|summary|certifications|languages))",
        "summary": r"(?:summary|profile|about|objective|career objective)\s*\n(.*?)(?=\n\s*\n|\n(?:experience|education|skills|projects|certifications|languages|contact))",
        "experience": r"(?:experience|work history|employment|professional experience)\s*\n(.*?)(?=\n\s*\n|\n(?:education|skills|projects|summary|certifications|languages|contact))",
        "education": r"(?:education|academic|qualifications)\s*\n(.*?)(?=\n\s*\n|\n(?:experience|skills|projects|summary|certifications|languages|contact))",
        "skills": r"(?:skills|technologies|technical skills|competencies)\s*\n(.*?)(?=\n\s*\n|\n(?:experience|education|projects|summary|certifications|languages|contact))",
        "projects": r"(?:projects|portfolio)\s*\n(.*?)(?=\n\s*\n|\n(?:experience|education|skills|summary|certifications|languages|contact))",
        "certifications": r"(?:certifications|certificates|licenses)\s*\n(.*?)(?=\n\s*\n|\n(?:experience|education|skills|projects|summary|languages|contact))",
        "languages": r"(?:languages|language skills)\s*\n(.*?)(?=\n\s*\n|\n(?:experience|education|skills|projects|summary|certifications|contact))",
    }
    sections = {}
    for name, pattern in section_patterns.items():
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            sections[name] = match.group(1).strip()
    return sections


def extract_email(text: str) -> Optional[str]:
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    match = re.search(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}", text)
    return match.group(0) if match else None


def extract_name_from_text(text: str) -> Optional[str]:
    lines = text.strip().split("\n")
    for line in lines[:5]:
        line = line.strip()
        if not line:
            continue
        if len(line.split()) == 2 and all(word[0].isupper() for word in line.split() if word):
            if not re.search(r"@|\.com|http|phone|email", line, re.IGNORECASE):
                return line
    return None
